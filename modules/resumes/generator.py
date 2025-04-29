"""
create_resume_docx.py
---------------------

Drop this file anywhere in your project and call create_resume_docx(...)
with your own data.  The function:

• creates resume/new_resume.docx   (using python-docx)
• creates resume/new_resume.pdf    (using fpdf2)
• auto-searches for DejaVuSans*.ttf anywhere under this folder
  – downloads the fonts if missing
  – falls back to built-in Helvetica and strips accents if it still
    can’t find a Unicode font
• never raises “Not enough horizontal space to render a single character”
  because every text block wraps inside an explicit usable width.
"""

from pathlib import Path
from urllib.request import urlretrieve, URLError
import unicodedata

import docx
from docx.shared import Pt, Inches
from fpdf import FPDF


# ---------------------------------------------------------------------
def create_resume_docx(
    user_details: dict,
    summary: str,
    experience: list[dict],
    projects: list[dict],
    skills: list[str],
    certificates: list[dict],
    max_bullets_per_section: int = 3,
):
    # =============================== DOCX =============================
    doc = docx.Document()

    # 0.5-inch margins
    sec = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Inches(0.5))

    # base fonts
    doc.styles["Heading 1"].font.size, doc.styles["Heading 1"].font.bold = Pt(12), True
    doc.styles["Heading 2"].font.size, doc.styles["Heading 2"].font.bold = Pt(11), True
    doc.styles["Normal"].font.size = Pt(10)

    # bullet paragraph style
    bullet_style = doc.styles.add_style("BulletTiny", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.paragraph_format.left_indent = Inches(0.2)
    bullet_style.paragraph_format.space_after = Pt(0)
    bullet_style.font.size = Pt(10)

    # header
    doc.add_heading(user_details["name"], 0)
    doc.paragraphs[-1].runs[0].font.size = Pt(14)
    doc.add_paragraph(
        f"{user_details['email']} | {user_details['phone_number']} | {user_details['address']}"
    )

    # summary
    doc.add_heading("SUMMARY", 1)
    doc.add_paragraph(summary, style="Normal")

    # experience
    doc.add_heading("EXPERIENCE", 1)
    for role in experience:
        doc.add_heading(role["company"], 2)
        doc.add_paragraph(f"{role['role']} | {role['dates']}", style="Normal")
        for i, line in enumerate(role["achievements"].splitlines()):
            if i >= max_bullets_per_section:
                break
            if line.strip():
                doc.add_paragraph(f"• {line.strip()}", style="BulletTiny")

    # projects
    doc.add_heading("PROJECTS", 1)
    for p in projects[:max_bullets_per_section]:
        doc.add_heading(p["name"], 2)
        doc.add_paragraph(f"{p['description']} | {p['technologies']}", style="Normal")

    # skills
    doc.add_heading("SKILLS", 1)
    doc.add_paragraph(", ".join(skills), style="Normal")

    # certificates
    if certificates:
        doc.add_heading("CERTIFICATES", 1)
        for c in certificates[:max_bullets_per_section]:
            doc.add_heading(c["name"], 2)
            doc.add_paragraph(c["description"], style="Normal")

    # save DOCX
    out_dir = Path("resume")
    out_dir.mkdir(exist_ok=True)
    docx_path = out_dir / "new_resume.docx"
    doc.save(docx_path)
    print("✔ DOCX saved →", docx_path)

    # =============================== PDF ==============================
    pdf = FPDF()
    pdf.set_margins(15, 12, 15)          # ~0.4 inch margins
    pdf.set_auto_page_break(False)
    pdf.add_page()

    # ---- locate or download DejaVu fonts ---------------------------
    def try_download(url: str, dest: Path):
        try:
            urlretrieve(url, dest)
        except URLError:
            return None
        return dest

    base_dir = Path(__file__).parent
    ttf_files = {p.name.lower(): p for p in base_dir.rglob("dejavusans*.ttf")}

    reg_ttf  = ttf_files.get("dejavusans.ttf")
    bold_ttf = ttf_files.get("dejavusans-bold.ttf")

    if not reg_ttf:
        reg_ttf = try_download(
            "https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/DejaVuSans.ttf",
            base_dir / "DejaVuSans.ttf",
        )
    if reg_ttf and not bold_ttf:
        bold_ttf = try_download(
            "https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/DejaVuSans-Bold.ttf",
            base_dir / "DejaVuSans-Bold.ttf",
        )

    if reg_ttf and reg_ttf.exists():
        pdf.add_font("U", "", str(reg_ttf),  uni=True)
        pdf.add_font("U", "B", str(bold_ttf or reg_ttf), uni=True)
        family = "U"
        bullet_char = "• "
        strip = lambda s: s
    else:
        print("⚠ Unicode font unavailable – using Helvetica & stripping accents")
        family = "Arial"
        bullet_char = "– "
        strip = lambda s: unicodedata.normalize("NFKD", s)\
                          .encode("latin-1", "ignore")\
                          .decode("latin-1")

    LEFT  = pdf.l_margin
    WIDTH = pdf.w - pdf.l_margin - pdf.r_margin

    # ---- text helpers -------------------------------------------------
    def H1(txt):
        pdf.set_font(family, "B", 12)
        pdf.set_x(LEFT)
        pdf.multi_cell(WIDTH, 6, strip(txt))
        pdf.ln(1)

    def H2(txt):
        pdf.set_font(family, "B", 11)
        pdf.set_x(LEFT + 2)
        pdf.multi_cell(WIDTH - 2, 5, strip(txt))

    def body(txt, indent=2, ln=0.5):
        pdf.set_font(family, "", 10)
        pdf.set_x(LEFT + indent)
        pdf.multi_cell(WIDTH - indent, 4.5, strip(txt))
        if ln: pdf.ln(ln)

    # ---- HEADER -------------------------------------------------------
    pdf.set_font(family, "B", 16)
    pdf.multi_cell(WIDTH, 8, strip(user_details["name"]), align="C")
    pdf.set_font(family, "", 9)
    header_line = f"{user_details['email']} | {user_details['phone_number']} | {user_details['address']}"
    pdf.multi_cell(WIDTH, 4, strip(header_line), align="C")
    pdf.ln(4)

    # ---- SUMMARY ------------------------------------------------------
    H1("SUMMARY")
    body(summary, indent=0); pdf.ln(2)

    # ---- EXPERIENCE ---------------------------------------------------
    H1("EXPERIENCE")
    for role in experience:
        H2(role["company"])
        body(f"{role['role']} | {role['dates']}", indent=4, ln=0.8)
        for i, line in enumerate(role["achievements"].splitlines()):
            if i >= max_bullets_per_section or not line.strip():
                break
            body(bullet_char + line.strip(), indent=6, ln=0.3)
        pdf.ln(1)

    # ---- PROJECTS -----------------------------------------------------
    H1("PROJECTS")
    for p in projects[:max_bullets_per_section]:
        H2(p["name"])
        body(f"{p['description']} | {p['technologies']}", indent=4)
        pdf.ln(1)

    # ---- SKILLS -------------------------------------------------------
    H1("SKILLS")
    body(", ".join(skills), indent=0)
    pdf.ln(2)

    # ---- CERTIFICATES -------------------------------------------------
    if certificates:
        H1("CERTIFICATES")
        for c in certificates[:max_bullets_per_section]:
            H2(c["name"])
            body(c["description"], indent=4)
            pdf.ln(1)

    # save PDF
    pdf_path = out_dir / "new_resume.pdf"
    pdf.output(pdf_path)
    print("✔ PDF  saved →", pdf_path)